"""Export helpers for PDF, DOCX, and TXT deliverables."""

from __future__ import annotations

from html import escape
from io import BytesIO
from typing import Iterable

from .schemas import GeneratedResult


def _safe_filename(value: str) -> str:
    allowed = [char.lower() if char.isalnum() else "-" for char in value]
    return "-".join("".join(allowed).split("-"))[:80] or "business-output"


def build_export_filename(result: GeneratedResult, extension: str) -> str:
    stem = _safe_filename(f"{result.module_name}-{result.title}")
    return f"{stem}.{extension.lstrip('.')}"


def result_to_markdown(result: GeneratedResult) -> str:
    lines = [
        f"# {result.title}",
        "",
        f"Module: {result.module_name}",
        f"Generated: {result.generated_at.strftime('%Y-%m-%d %H:%M UTC')}",
        f"Model: {result.model}",
        "",
    ]
    for section, content in result.sections.items():
        lines.extend([f"## {section}", "", content.strip(), ""])
    return "\n".join(lines).strip() + "\n"


def export_txt(result: GeneratedResult) -> bytes:
    return result_to_markdown(result).encode("utf-8")


def _paragraph_lines(text: str) -> Iterable[str]:
    for line in text.splitlines():
        yield line.strip() if line.strip() else " "


def export_docx(result: GeneratedResult) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    document.add_heading(result.title, level=0)
    meta = document.add_paragraph()
    meta.add_run(f"{result.module_name} | {result.generated_at.strftime('%Y-%m-%d %H:%M UTC')} | {result.model}")

    for section_name, content in result.sections.items():
        document.add_heading(section_name, level=1)
        for line in _paragraph_lines(content):
            if line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_pdf(result: GeneratedResult) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "EnterpriseTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#111827"),
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        textColor=colors.HexColor("#4b5563"),
        spaceAfter=16,
    )
    heading_style = ParagraphStyle(
        "EnterpriseHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#0f766e"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "EnterpriseBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#1f2937"),
    )

    story = [
        Paragraph(escape(result.title), title_style),
        Paragraph(
            escape(f"{result.module_name} | {result.generated_at.strftime('%Y-%m-%d %H:%M UTC')} | {result.model}"),
            meta_style,
        ),
    ]

    for section_name, content in result.sections.items():
        story.append(Paragraph(escape(section_name), heading_style))
        for line in _paragraph_lines(content):
            prefix = "&#8226; " if line.startswith("- ") else ""
            clean_line = line[2:] if line.startswith("- ") else line
            story.append(Paragraph(prefix + escape(clean_line), body_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    return buffer.getvalue()
